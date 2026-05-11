"""
Educator Copilot — DOCX Exporter
==================================
Exports generated RPS to DOCX based on institution preset.
"""

import io
import json
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

from tools.database import get_connection, load_rps_weeks
from lib.template_presets import get_preset

def export_rps_to_docx(course_id: str, preset_id: str) -> bytes:
    """Generate DOCX format RPS."""
    if Document is None:
        raise RuntimeError("python-docx belum terinstal. Jalankan: pip install python-docx")

    conn = get_connection()
    course = conn.execute("SELECT * FROM courses WHERE course_id = ?", (course_id,)).fetchone()
    conn.close()

    if not course:
        raise ValueError("Course not found")

    course = dict(course)
    preset = get_preset(preset_id)
    weeks = load_rps_weeks(course_id)

    doc = Document()
    
    # Title
    heading = doc.add_heading(f"RENCANA PEMBELAJARAN SEMESTER (RPS)", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if preset.has_rpmk:
         doc.add_paragraph("*) Dokumen ini juga mencakup RPMK sesuai ketentuan ITB").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("I. IDENTITAS MATA KULIAH", level=1)
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    def add_row(k, v):
        row_cells = table.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = str(v)

    add_row("Nama Mata Kuliah", course['name'])
    if course.get('name_en'):
        add_row("Nama MK (EN)", course['name_en'])
    add_row("Kode Mata Kuliah", course['code'])
    add_row("SKS", f"{course['credits']} SKS")
    add_row("Deskripsi", course.get('description', ''))
    
    # Adding more detailed info based on form_data
    form_data = {}
    if course.get("form_data"):
        try:
             form_data = json.loads(course["form_data"])
        except Exception:
             pass

    if form_data.get("rumpun_mk"):
         add_row("Rumpun MK", form_data["rumpun_mk"])
    if form_data.get("jenis_mk"):
         add_row("Jenis MK", form_data["jenis_mk"])
         
    doc.add_heading("II. JADWAL MINGGUAN", level=1)
    
    # Constructing a comprehensive table
    headers = ["Mg", "Sub-CPMK", "Bahan Kajian/Topik", "Metode"]
    if preset.has_pt_bm:
        headers.extend(["TM", "PT", "BM"])
    headers.extend(["Pengalaman Belajar"])
    
    table2 = doc.add_table(rows=1, cols=len(headers))
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        
    for w in weeks:
        row_cells = table2.add_row().cells
        row_cells[0].text = str(w['week_number'])
        
        # Sub-cpmk
        sc = w.get('sub_cpmk', '')
        if isinstance(sc, list): sc = ", ".join(sc)
        row_cells[1].text = str(sc)
        
        # Topik
        top = w.get('topics', [])
        if isinstance(top, list): top = ", ".join(top)
        row_cells[2].text = str(top)
        
        # Metode
        row_cells[3].text = w.get('teaching_method', '')
        
        col_idx = 4
        if preset.has_pt_bm:
             row_cells[col_idx].text = str(w.get('time_tm', ''))
             row_cells[col_idx+1].text = str(w.get('time_pt', ''))
             row_cells[col_idx+2].text = str(w.get('time_bm', ''))
             col_idx += 3
             
        row_cells[col_idx].text = w.get('student_activity', '')
        
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream.read()
